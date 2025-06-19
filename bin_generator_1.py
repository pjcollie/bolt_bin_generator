import tkinter as tk
from tkinter import messagebox
import re
from collections import defaultdict
from docx import Document
from docx.shared import Inches

class BoltBinApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Bolt Bin Layout Generator")
        self.root.geometry("800x600")

        self.bolts = []  # List to store (size, size_display, length, length_num) tuples
        self.bin_layout = []
        self.rows, self.cols = 0, 0

        # Thread pitch mapping for coarse threads
        self.thread_pitches = {
            "#6": "32", "#8": "32", "#10": "24",
            "1/4": "20", "5/16": "18", "3/8": "16", "1/2": "13"
        }

        self.create_widgets()

    def create_widgets(self):
        # Bin Size Selection
        tk.Label(self.root, text="Select Bin Size:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.bin_size = tk.StringVar(value="56")
        tk.Radiobutton(self.root, text="56 Slots (7x8)", variable=self.bin_size, value="56").grid(row=0, column=1, sticky="w")
        tk.Radiobutton(self.root, text="72 Slots (8x9)", variable=self.bin_size, value="72").grid(row=0, column=2, sticky="w")

        # Material (fixed to Grade 5 Zinc)
        tk.Label(self.root, text="Material: Grade 5 Zinc").grid(row=1, column=0, padx=5, pady=5, sticky="e")

        # Bolt Size Input
        tk.Label(self.root, text="Bolt Size (e.g., #6, 1/4, 1-1/2):").grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.size_entry = tk.Entry(self.root)
        self.size_entry.grid(row=2, column=1, columnspan=2, sticky="w")

        # Length Input
        tk.Label(self.root, text="Length in inches (e.g., 1, 1/2, 1-1/2):").grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.length_entry = tk.Entry(self.root)
        self.length_entry.grid(row=3, column=1, columnspan=2, sticky="w")

        # Add Bolt Button
        tk.Button(self.root, text="Add Bolt", command=self.add_bolt).grid(row=4, column=1, pady=10)

        # Bolt List Display
        tk.Label(self.root, text="Added Bolts:").grid(row=5, column=0, padx=5, pady=5, sticky="ne")
        self.bolt_listbox = tk.Listbox(self.root, height=10, width=50)
        self.bolt_listbox.grid(row=5, column=1, columnspan=2, padx=5, pady=5)

        # Generate Layout Button
        tk.Button(self.root, text="Generate Layout", command=self.generate_layout).grid(row=6, column=1, pady=10)

        # Save Layout Button
        tk.Button(self.root, text="Save Layout", command=self.save_layout).grid(row=6, column=2, pady=10)

        # Layout Display Frame
        self.layout_frame = tk.Frame(self.root)
        self.layout_frame.grid(row=7, column=0, columnspan=3, padx=5, pady=5)

    def parse_fraction(self, value):
        """Convert fraction, mixed fraction, or decimal to float (e.g., '1/2' -> 0.5, '1-1/2' -> 1.5)."""
        try:
            # Handle mixed fractions (e.g., "1-1/2")
            if '-' in value:
                whole, frac = value.split('-')
                whole_num = int(whole)
                if '/' in frac:
                    num, denom = map(int, frac.split('/'))
                    if denom == 0:
                        return None
                    return whole_num + num / denom
                return None
            # Handle simple fractions (e.g., "1/2")
            if '/' in value:
                num, denom = map(int, value.split('/'))
                if denom == 0:
                    return None
                return num / denom
            # Handle decimals (e.g., "0.5")
            return float(value)
        except (ValueError, ZeroDivisionError):
            return None

    def add_bolt(self):
        size = self.size_entry.get().strip()
        length = self.length_entry.get().strip()

        # Validate inputs
        if not size or not length:
            messagebox.showerror("Error", "Size and length are required!")
            return

        # Validate size (e.g., #6, 1/4, 1-1/2)
        size_valid = re.match(r"^(#?\d+|\d+/\d+|\d+-\d+/\d+)$", size)
        if not size_valid:
            messagebox.showerror("Error", "Invalid size format! Use e.g., #6, 1/4, 1-1/2")
            return

        # Parse length (fraction, mixed fraction, or decimal)
        length_num = self.parse_fraction(length)
        if length_num is None:
            messagebox.showerror("Error", "Invalid length format! Use e.g., 1, 1/2, 1-1/2")
            return

        # Get thread pitch
        thread_pitch = self.thread_pitches.get(size, "Coarse")
        size_display = f"{size}-{thread_pitch}" if thread_pitch != "Coarse" else f"{size}-Coarse"

        # Add to bolts list
        self.bolts.append((size, size_display, length, length_num))
        self.bolt_listbox.insert(tk.END, f"Grade 5 Zinc {size_display} x {length}")
        self.size_entry.delete(0, tk.END)
        self.length_entry.delete(0, tk.END)

    def parse_size(self, size):
        """Extract numeric value for sorting (e.g., '#6' -> 6, '1/4' -> 0.25, '1-1/2' -> 1.5)."""
        if size.startswith('#'):
            match = re.match(r"#(\d+)", size)
            if match:
                return float(match.group(1))
        return self.parse_fraction(size) or size  # Fallback to string if invalid

    def generate_layout(self):
        # Set grid dimensions
        if self.bin_size.get() == "56":
            self.rows, self.cols = 7, 8
        else:
            self.rows, self.cols = 8, 9

        # Group bolts by size
        bolts_by_size = defaultdict(list)
        for bolt in self.bolts:
            size, size_display, length, length_num = bolt
            bolts_by_size[size].append((size_display, length, length_num))

        # Check if too many unique sizes
        unique_sizes = len(bolts_by_size)
        if unique_sizes > self.rows:
            messagebox.showerror("Error", f"Too many unique sizes ({unique_sizes}) for {self.rows} rows!")
            return

        # Check if any size has too many bolts for a row
        for size, bolts in bolts_by_size.items():
            if len(bolts) > self.cols:
                messagebox.showerror("Error", f"Size {size} has too many bolts ({len(bolts)}) for {self.cols} columns!")
                return

        # Sort sizes by numeric value
        sorted_sizes = sorted(bolts_by_size.keys(), key=self.parse_size)

        # Initialize grid
        self.bin_layout = [["Empty" for _ in range(self.cols)] for _ in range(self.rows)]

        # Fill grid: one size per row, lengths sorted left to right
        for row, size in enumerate(sorted_sizes):
            # Sort bolts of this size by length
            sorted_bolts = sorted(bolts_by_size[size], key=lambda x: x[2])
            for col, (size_display, length, _) in enumerate(sorted_bolts):
                self.bin_layout[row][col] = f"Grade 5 Zinc {size_display} x {length}"

        # Clear previous layout
        for widget in self.layout_frame.winfo_children():
            widget.destroy()

        # Display grid
        for row in range(self.rows):
            for col in range(self.cols):
                tk.Label(self.layout_frame, text=self.bin_layout[row][col], borderwidth=1, relief="solid", width=20, height=2).grid(row=row, column=col, padx=1, pady=1)

        # Print to console
        print("\nBolt Bin Layout:")
        for row in self.bin_layout:
            print(" | ".join(f"{cell:<20}" for cell in row))

    def save_layout(self):
        if not self.bin_layout:
            messagebox.showerror("Error", "Generate a layout first!")
            return

        try:
            # Create Word document
            doc = Document()
            doc.add_heading(f"Bolt Bin Layout ({self.rows}x{self.cols})", level=1)

            # Add table
            table = doc.add_table(rows=self.rows, cols=self.cols)
            table.style = 'Table Grid'  # Add borders
            for row in range(self.rows):
                for col in range(self.cols):
                    cell = table.cell(row, col)
                    cell.text = self.bin_layout[row][col]
                    # Set cell width (approx. 1 inch per column)
                    cell.width = Inches(1.0)

            # Save document
            doc.save("bolt_bin_layout.docx")
            messagebox.showinfo("Success", "Layout saved to bolt_bin_layout.docx")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save document: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = BoltBinApp(root)
    root.mainloop()
    